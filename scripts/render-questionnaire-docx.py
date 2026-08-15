#!/usr/bin/env python3
"""Render a locale- and version-specific readiness catalog to DOCX."""

import argparse
import json
import os
import subprocess
import tempfile
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree as ET


COPY = {
    "ko": {
        "title": "글로벌 진출 준비도 진단 설문",
        "subtitle": "창업자 작성용 · {count}문항 · v{version}",
        "info": ["회사명", "작성자 · 직책", "작성일", "목표 국가", "주요 제품·서비스"],
        "guide": "작성 안내",
        "instructions": [
            "각 문항은 지금 상태에 가장 가까운 보기 하나를 고르시면 됩니다. 아직 하지 않은 것을 고르셔도 불이익은 없습니다.",
            "①은 ‘그 부분까지 생각해보지 못했다’, ②는 ‘필요한 줄은 알지만 아직 못 했다’, ③은 ‘실제로 해봤다’, ④는 ‘반복됐거나 외부에서 확인받았다’는 뜻입니다.",
            "③이나 ④를 고르신 문항만 아래 칸에 간단히 적어주세요. ①·②를 고르신 문항은 비워두셔도 됩니다.",
            "계약서, 고객명부, 재무자료 원본은 제출하지 않으셔도 됩니다. 고객사 이름은 ‘고객 A’처럼 익명으로 적으셔도 됩니다.",
            "각 단계의 정규화된 배점 기준 80% 이상이 ③ 또는 ④이고 미해결 Critical 선결 조건이 없으면 그 단계를 통과합니다.",
        ],
        "follow_up": "③·④를 고르셨다면",
        "done": "문항",
    },
    "en": {
        "title": "Global Market Entry Readiness Assessment",
        "subtitle": "Founder Workbook · {count} Questions · v{version}",
        "info": ["Company", "Founder · Role", "Date", "Target Country", "Primary Product or Service"],
        "guide": "How to Complete This Assessment",
        "instructions": [
            "Choose the option that best reflects where you are today. There is no penalty for selecting work you have not started.",
            "① means ‘not considered yet,’ ② ‘recognized or planned,’ ③ ‘executed with an example,’ and ④ ‘repeated or externally validated.’",
            "For answers ③ or ④, add a brief note in the space provided. You may leave the space blank for answers ① and ②.",
            "Do not submit original contracts, customer lists, or financial records. You may anonymize customer names, for example as ‘Customer A.’",
            "A stage passes when at least 80% of its weighted score is rated ③ or ④ and no critical prerequisite remains unresolved.",
        ],
        "follow_up": "If you selected ③ or ④",
        "done": "questions",
    },
}
FONT = "Noto Serif KR"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def find_korean_font():
    configured = os.environ.get("READINESS_DOCX_FONT")
    if configured:
        font = Path(configured).expanduser()
        if font.is_file():
            return font
        raise FileNotFoundError(f"READINESS_DOCX_FONT does not exist: {font}")

    resource_root = Path.home() / "Documents" / "Developer Resource"
    font = next(resource_root.rglob("NotoSerifKR-Regular.ttf"), None) if resource_root.exists() else None
    if font:
        return font
    raise FileNotFoundError("Set READINESS_DOCX_FONT to an embeddable Korean TTF font.")


def embed_korean_font(docx_path):
    font_path = find_korean_font()
    font_key = uuid.uuid4()
    font_bytes = bytearray(font_path.read_bytes())
    # ECMA-376 uses the GUID bytes in reverse order for the 32-byte XOR prefix.
    key = font_key.bytes[::-1]
    for index in range(min(32, len(font_bytes))):
        font_bytes[index] ^= key[index % len(key)]

    with zipfile.ZipFile(docx_path) as source:
        entries = {item.filename: (item, source.read(item.filename)) for item in source.infolist()}

    font_table = ET.fromstring(entries["word/fontTable.xml"][1])
    font_node = font_table.find(f".//{{{WORD_NS}}}font[@{{{WORD_NS}}}name='{FONT}']")
    if font_node is None:
        font_node = ET.SubElement(font_table, f"{{{WORD_NS}}}font", {f"{{{WORD_NS}}}name": FONT})
    for child in list(font_node):
        if child.tag == f"{{{WORD_NS}}}embedRegular":
            font_node.remove(child)
    ET.SubElement(
        font_node,
        f"{{{WORD_NS}}}embedRegular",
        {f"{{{REL_NS}}}id": "rIdEmbeddedKoreanFont", f"{{{WORD_NS}}}fontKey": f"{{{str(font_key).upper()}}}"},
    )
    entries["word/fontTable.xml"] = (
        entries["word/fontTable.xml"][0],
        ET.tostring(font_table, encoding="utf-8", xml_declaration=True),
    )

    rels_path = "word/_rels/fontTable.xml.rels"
    rels = ET.fromstring(entries[rels_path][1]) if rels_path in entries else ET.Element(f"{{{PACKAGE_REL_NS}}}Relationships")
    for child in list(rels):
        if child.attrib.get("Id") == "rIdEmbeddedKoreanFont":
            rels.remove(child)
    ET.SubElement(
        rels,
        f"{{{PACKAGE_REL_NS}}}Relationship",
        {
            "Id": "rIdEmbeddedKoreanFont",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
            "Target": "fonts/NotoSerifKR-Regular.odttf",
        },
    )
    rels_bytes = ET.tostring(rels, encoding="utf-8", xml_declaration=True)

    content_types = ET.fromstring(entries["[Content_Types].xml"][1])
    if not any(child.attrib.get("Extension") == "odttf" for child in content_types):
        ET.SubElement(
            content_types,
            f"{{{CONTENT_TYPE_NS}}}Default",
            {
                "Extension": "odttf",
                "ContentType": "application/vnd.openxmlformats-officedocument.obfuscatedFont",
            },
        )
    entries["[Content_Types].xml"] = (
        entries["[Content_Types].xml"][0],
        ET.tostring(content_types, encoding="utf-8", xml_declaration=True),
    )

    with tempfile.NamedTemporaryFile(dir=docx_path.parent, suffix=".docx", delete=False) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for name, (item, data) in entries.items():
                target.writestr(item, data)
            target.writestr(rels_path, rels_bytes)
            target.writestr("word/fonts/NotoSerifKR-Regular.odttf", font_bytes)
        os.replace(temp_path, docx_path)
    finally:
        temp_path.unlink(missing_ok=True)


def set_cell_shading(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, value=100):
    margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell._tc.get_or_add_tcPr().append(margins)
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def add_text(paragraph, text, *, size=10, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if any("가" <= char <= "힣" for char in text):
        language = OxmlElement("w:lang")
        language.set(qn("w:val"), "ko-KR")
        language.set(qn("w:eastAsia"), "ko-KR")
        run._element.rPr.append(language)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_paragraph(doc, text="", *, size=10, bold=False, italic=False, color=None, before=0, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    add_text(paragraph, text, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def load_catalog(root, locale, version, node):
    command = [node, str(root / "scripts/build-questionnaire-docx.js"), "--json", "--locale", locale, "--version", version]
    return json.loads(subprocess.run(command, cwd=root, check=True, capture_output=True, text=True).stdout)


def render(output, locale, version, node):
    root = Path(__file__).resolve().parents[1]
    catalog = load_catalog(root, locale, version, node)
    copy = COPY[locale]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    add_paragraph(doc, copy["title"], size=22, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, copy["subtitle"].format(count=len(catalog["questions"]), version=version), color="666666", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    info = doc.add_table(rows=0, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    for label in copy["info"]:
        row = info.add_row()
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(5.5)
        set_cell_shading(row.cells[0], "F2F2F2")
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(row.cells[0].paragraphs[0], label, bold=True)
        add_text(row.cells[1].paragraphs[0], "")

    add_paragraph(doc, copy["guide"], size=13, bold=True, before=14, after=6)
    for instruction in copy["instructions"]:
        add_paragraph(doc, f"• {instruction}", after=3)

    items_by_stage = {}
    for item in catalog["items"]:
        items_by_stage.setdefault(item["stageId"], []).append(item)
    questions_by_item = {}
    for question in catalog["questions"]:
        questions_by_item.setdefault(question["itemId"], []).append(question)

    q_no = 0
    for stage in catalog["stages"]:
        heading = add_paragraph(doc, stage["label"], size=17, bold=True, after=4)
        heading.paragraph_format.keep_with_next = True
        add_paragraph(doc, stage["intro"], color="666666", after=12)
        for item in items_by_stage.get(stage["id"], []):
            keep_with_next(add_paragraph(doc, item["label"], size=13, bold=True, before=8, after=5))
            for question in questions_by_item.get(item["id"], []):
                q_no += 1
                paragraph = keep_with_next(add_paragraph(doc, before=7, after=4))
                critical = "★ " if question.get("critical") else ""
                add_text(paragraph, f"Q{q_no}. {critical}", size=10.5, bold=True)
                add_text(paragraph, question["question"], size=10.5)
                if question.get("help"):
                    keep_with_next(add_paragraph(doc, question["help"], size=9, italic=True, color="666666", after=3))
                for index, option in enumerate(question["options"]):
                    option_paragraph = keep_with_next(add_paragraph(doc, f"☐ {'①②③④'[index]} {option}", after=2))
                    option_paragraph.paragraph_format.left_indent = Inches(0.2)
                keep_with_next(add_paragraph(doc, f'{copy["follow_up"]} — {question["followUp"]}', size=9, italic=True, color="666666", before=3, after=3))
                box = doc.add_table(rows=1, cols=1)
                box.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_cell_margins(box.cell(0, 0), 120)
                add_text(box.cell(0, 0).paragraphs[0], "\n")
                cant_split = OxmlElement("w:cantSplit")
                box.rows[0]._tr.get_or_add_trPr().append(cant_split)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    if locale == "ko":
        embed_korean_font(output)
    print(f"{output} — {q_no} {copy['done']}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    parser.add_argument("locale", choices=("ko", "en"))
    parser.add_argument("--version", choices=("4.0", "5.0"), default="4.0")
    parser.add_argument("--node", default="node")
    args = parser.parse_args()
    render(args.output, args.locale, args.version, args.node)
